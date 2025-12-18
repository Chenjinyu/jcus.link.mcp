"""Document parser service for handling PDF, HTML, and DOC files."""

import os
import tempfile
from pathlib import Path
from typing import Optional
import logging

from pypdf import PdfReader
from docx import Document as DocxDocument
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse various document formats and extract text content."""

    @staticmethod
    async def parse_file(file_bytes: bytes, filename: str, content_type: str) -> str:
        """
        Parse a file and extract text content.

        Args:
            file_bytes: The file content as bytes
            filename: Original filename
            content_type: MIME type of the file

        Returns:
            Extracted text content

        Raises:
            ValueError: If file type is not supported
            Exception: If parsing fails
        """
        file_extension = Path(filename).suffix.lower()

        # Save to temporary file for parsing
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
            tmp_file.write(file_bytes)
            tmp_path = tmp_file.name

        try:
            if file_extension == ".pdf":
                return await DocumentParser._parse_pdf(tmp_path)
            elif file_extension in [".doc", ".docx"]:
                return await DocumentParser._parse_docx(tmp_path)
            elif file_extension in [".html", ".htm"]:
                return await DocumentParser._parse_html(file_bytes)
            elif file_extension == ".txt":
                return file_bytes.decode("utf-8", errors="replace")
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")

        finally:
            # Clean up temporary file
            if os.path.exists(tmp_path):
                os.remove(tmp_path)

    @staticmethod
    async def _parse_pdf(file_path: str) -> str:
        """Parse PDF file and extract text."""
        try:
            reader = PdfReader(file_path)
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return "\n\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            raise Exception(f"Failed to parse PDF: {str(e)}")

    @staticmethod
    async def _parse_docx(file_path: str) -> str:
        """Parse DOCX file and extract text."""
        try:
            doc = DocxDocument(file_path)
            text_parts = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            raise Exception(f"Failed to parse DOCX: {str(e)}")

    @staticmethod
    async def _parse_html(html_bytes: bytes) -> str:
        """Parse HTML file and extract text content."""
        try:
            soup = BeautifulSoup(html_bytes, "lxml")

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Get text content
            text = soup.get_text(separator="\n", strip=True)

            # Clean up excessive whitespace
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            return "\n".join(lines)

        except Exception as e:
            logger.error(f"Error parsing HTML: {e}")
            raise Exception(f"Failed to parse HTML: {str(e)}")

