# src/utils/document_parser.py
"""
Document parser for various file formats (PDF, DOCX, HTML, TXT, URL)
"""

import logging
import requests
import validators
from pathlib import Path
from typing import Optional, Union
from io import BytesIO

# Document parsing libraries
from PyPDF2 import PdfReader
from docx import Document
from bs4 import BeautifulSoup
import markdown

from ..config import settings

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse documents from various sources and formats"""
    
    @staticmethod
    async def parse(
        content: Union[bytes, str],
        file_type: Optional[str] = None,
        is_url: bool = False
    ) -> str:
        """
        Parse document content and extract text
        
        Args:
            content: File content (bytes) or URL (str)
            file_type: File extension (.pdf, .docx, etc.)
            is_url: Whether content is a URL
            
        Returns:
            Extracted text content
        """
        try:
            if is_url:
                return await DocumentParser._parse_url(str(content))
            
            # Ensure content is bytes for file parsing
            if isinstance(content, str):
                content = content.encode('utf-8')
            
            if file_type == ".pdf":
                return DocumentParser._parse_pdf(content)
            elif file_type in [".doc", ".docx"]:
                return DocumentParser._parse_docx(content)
            elif file_type == ".html":
                return DocumentParser._parse_html(content)
            elif file_type == ".md":
                return DocumentParser._parse_markdown(content)
            elif file_type == ".txt":
                return DocumentParser._parse_txt(content)
            else:
                # Try to parse as text by default
                return DocumentParser._parse_txt(content)
        
        except Exception as e:
            logger.error(f"Document parsing failed: {e}")
            raise ValueError(f"Failed to parse document: {str(e)}")
    
    @staticmethod
    def _parse_pdf(content: bytes) -> str:
        """Parse PDF file"""
        try:
            pdf_file = BytesIO(content)
            pdf_reader = PdfReader(pdf_file)
            
            # Limit pages for performance
            max_pages = min(
                len(pdf_reader.pages),
                settings.pdf_max_pages
            )
            
            text_parts = []
            for page_num in range(max_pages):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            full_text = "\n\n".join(text_parts)
            
            if not full_text.strip():
                raise ValueError("No text content extracted from PDF")
            
            logger.info(f"Parsed PDF: {len(full_text)} characters from {max_pages} pages")
            return full_text
        
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    @staticmethod
    def _parse_docx(content: bytes) -> str:
        """Parse DOCX file"""
        try:
            docx_file = BytesIO(content)
            doc = Document(docx_file)
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            full_text = "\n\n".join(paragraphs)
            
            if not full_text.strip():
                raise ValueError("No text content extracted from DOCX")
            
            logger.info(f"Parsed DOCX: {len(full_text)} characters")
            return full_text
        
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    @staticmethod
    def _parse_html(content: Union[bytes, str]) -> str:
        """Parse HTML content"""
        try:
            if isinstance(content, bytes):
                content = content.decode('utf-8', errors='ignore')
            
            soup = BeautifulSoup(content, 'lxml')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            if not text.strip():
                raise ValueError("No text content extracted from HTML")
            
            logger.info(f"Parsed HTML: {len(text)} characters")
            return text
        
        except Exception as e:
            logger.error(f"HTML parsing error: {e}")
            raise ValueError(f"Failed to parse HTML: {str(e)}")
    
    @staticmethod
    def _parse_markdown(content: Union[bytes, str]) -> str:
        """Parse Markdown content"""
        try:
            if isinstance(content, bytes):
                content = content.decode('utf-8', errors='ignore')
            
            # Convert markdown to HTML first, then extract text
            html = markdown.markdown(content)
            return DocumentParser._parse_html(html)
        
        except Exception as e:
            logger.error(f"Markdown parsing error: {e}")
            raise ValueError(f"Failed to parse Markdown: {str(e)}")
    
    @staticmethod
    def _parse_txt(content: Union[bytes, str]) -> str:
        """Parse plain text file"""
        try:
            if isinstance(content, bytes):
                # Try different encodings
                for encoding in ['utf-8', 'latin-1', 'cp1252']:
                    try:
                        text = content.decode(encoding)
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    # Fallback: decode with errors='ignore'
                    text = content.decode('utf-8', errors='ignore')
            else:
                text = content
            
            if not text.strip():
                raise ValueError("Empty text file")
            
            logger.info(f"Parsed text: {len(text)} characters")
            return text
        
        except Exception as e:
            logger.error(f"Text parsing error: {e}")
            raise ValueError(f"Failed to parse text: {str(e)}")
    
    @staticmethod
    async def _parse_url(url: str) -> str:
        """Parse content from URL"""
        try:
            # Validate URL
            if not validators.url(url):
                raise ValueError(f"Invalid URL: {url}")
            
            # Check allowed schemes
            if not any(url.startswith(scheme) for scheme in settings.allowed_url_schemes):
                raise ValueError(
                    f"URL scheme not allowed. Must be one of: {settings.allowed_url_schemes}"
                )
            
            logger.info(f"Fetching URL: {url}")
            
            # Fetch URL content
            response = requests.get(
                url,
                timeout=settings.url_fetch_timeout,
                headers={'User-Agent': 'Mozilla/5.0 (Resume Parser Bot)'}
            )
            response.raise_for_status()
            
            # Determine content type
            content_type = response.headers.get('Content-Type', '').lower()
            
            if 'application/pdf' in content_type:
                return DocumentParser._parse_pdf(response.content)
            elif 'text/html' in content_type:
                return DocumentParser._parse_html(response.content)
            elif 'text/plain' in content_type:
                return DocumentParser._parse_txt(response.content)
            else:
                # Try to parse as HTML by default
                return DocumentParser._parse_html(response.content)
        
        except requests.RequestException as e:
            logger.error(f"URL fetch error: {e}")
            raise ValueError(f"Failed to fetch URL: {str(e)}")
        except Exception as e:
            logger.error(f"URL parsing error: {e}")
            raise ValueError(f"Failed to parse URL content: {str(e)}")
    
    @staticmethod
    def detect_file_type(filename: str) -> str:
        """Detect file type from filename"""
        suffix = Path(filename).suffix.lower()
        
        if suffix not in settings.allowed_file_types:
            raise ValueError(
                f"File type '{suffix}' not allowed. "
                f"Allowed types: {settings.allowed_file_types}"
            )
        
        return suffix
    
    @staticmethod
    def is_url(text: str) -> bool:
        """Check if text is a valid URL"""
        return validators.url(text) is True


# Singleton instance
_document_parser: Optional[DocumentParser] = None


def get_document_parser() -> DocumentParser:
    """Get document parser instance (singleton)"""
    global _document_parser
    
    if _document_parser is None:
        _document_parser = DocumentParser()
    
    return _document_parser